import time
import streamlit as st
import os
from io import StringIO
from calling import call_watsonx
from assemblerCode import split_assem_code, clean_assem_code
from javaCode import split_java_code
from cobolCode import split_cbl_code, process_cobol_file, split_off_procedure
from pythonCode import split_py_code
from basicCode import split_basic_code
from docx import Document
from datetime import datetime

class StreamlitApp:
    def __init__(self):
        st.set_page_config(layout="wide")
        self.chunk_dict = {}
        self.selected_code = ""  
        self.toggle_summarize = False
        self.toggle_explanation = False
        self.toggle_interactions = False
        self.toggle_walkthrough_not = False
        self.toggle_walkthrough_ie = False
        self.custom_question = "" 
        # Initialize session state variables
        st.session_state.context_valid = True     

    #when user uploads a file this takes it to the appropriate splitter
    def process_file_by_type(self, file, file_path):
        # Determine file type based on extension
        file_extension = file_path.split('.')[-1].lower()

        if file_extension == 'cbl':
            code_type = "cobol"
            #remove the prior to procedure division section
            file_content=split_off_procedure(file)
            #clean up the cobol code
            pretty_cobol =process_cobol_file(file_content)
            #split cobol into watson sized chunks
            cobol_chunks = split_cbl_code(pretty_cobol)
            return cobol_chunks, code_type
        elif file_extension == 'asm':
            code_type = "assembler"
            pretty_assem = clean_assem_code(file)
            assembler_chunks =split_assem_code(pretty_assem)
            return assembler_chunks, code_type
        elif file_extension == 'py':
            code_type = "python"
            python_chunks = split_py_code(file)
            return python_chunks, code_type
        elif file_extension == 'java':
            code_type = "java"
            java_chunks = split_java_code(file)
            return java_chunks, code_type
        elif file_extension == 'bas':
            code_type = "BASIC"
            basic_chunks = split_basic_code(file)
            return basic_chunks, code_type    
        else:
            print(f"Unsupported file type: {file_extension}")
            return []

    def columns_shift(self, chunks):
        # assembling the keys of the chunk dictionary
        dropdown_items = list(self.chunk_dict.keys())

        # Display the "Paste in my own code" option in the dropdown
        dropdown_items.insert(0, "Paste in my own code")
        self.selected_code = st.selectbox("Select a section of code:", dropdown_items, format_func=lambda x: x, index=0)

        # Add a toggle to switch between text_area and code display
        display_code = st.checkbox("Display formatted code")

        # Use st.empty() to create a placeholder
        code_display_placeholder = st.empty()

        # If the user selects "Paste in my own code", provide a text area for manual input
        if self.selected_code == "Paste in my own code":
            # Define custom CSS to style the warning notice
            custom_css = """
            <style>
                .custom-warning {
                    color: red;
                    font-family: halvetica; /* Replace with your desired font */
                    font-size: 14px; /* Adjust the font size as needed */
                }
            </style>
            """

            # Display the custom CSS to style the warning notice
            st.markdown(custom_css, unsafe_allow_html=True)

            # Display the text area for pasting code
            self.text_area_code = st.text_area("Paste your code here:", key="pasted_code", height=870)
                # Hide the "Code:" text area by clearing the placeholder
            code_display_placeholder.empty()

            # Display the "Character count of the section of code" notice with custom styling
            st.markdown('<p class="custom-warning">Code length must not exceed 6000 characters!</p>', unsafe_allow_html=True)
        else:
            # getting the user selection of the dropdown menu that is displayed
            self.text_area_code = self.chunk_dict.get(self.selected_code, "")

            # Display the selected chunk with new lines
            if dropdown_items:
                if display_code and self.selected_code != "Paste in my own code":
                    # Check if the selected code is not empty before using st.code
                    selected_code = self.chunk_dict.get(self.selected_code, "")
                    if selected_code:
                        # Display code using st.code
                        st.code(selected_code, language='cobol')
                else:
                    # Display code in text_area
                    st.text_area("Code:", value=self.text_area_code, height=900)       
        
        st.session_state['previous_selected_code'] = self.selected_code

    # user adds a file and nav bar is displayed, and drop down is populated
    def update_ui_after_file_upload(self, chunks, code_type):
        # Reset session state variables
        st.session_state.context_valid = True

        #assemble
        self.chunk_dict = {chunk.split('\n', 1)[0].strip(): chunk.strip() for chunk in chunks}

        # Create two columns
        col1, col2 = st.columns([1, 1])

        # Column 1: Code selection and display
        with col1:
            self.columns_shift(chunks)
            # SIDEBAR (left side)----->

            #Add the ability for the user to add context to the prompt.
            # Display the "Add Context" checkbox
            add_context = st.sidebar.checkbox("Add Context", help="Paste in any specific data for this program or company-specific verbiage/jargon.")
            st.session_state.add_context = add_context

            context = "No context this time."
            if add_context:
                context = st.sidebar.text_area("Paste in Context (max 850 characters):")
                context_valid = len(context) <= 850
                st.session_state.context_valid = context_valid

                if not context_valid:
                    st.sidebar.warning("Context must not exceed 850 characters.")                          
    
            # Display the length of the dictionary
            num_dict_items = len(self.chunk_dict)
            st.sidebar.text (" ")
            st.sidebar.write(f"Number of sections of code: {num_dict_items}")

            # how many characters in the section selected
            char_count = len(self.text_area_code)
            st.sidebar.write(f"Character count of the section of code: {char_count}")
            st.sidebar.text (" ")

            # Display the radio buttons for selecting the option
            selected_option = st.sidebar.radio("Select an analysis type:", ["Code Summarization", "Code Walkthrough", "Custom Question"])

            # Display the text area for a custom question if selected option is "Custom Question"
            custom_question = ""
            if selected_option == "Custom Question":
                custom_question = st.sidebar.text_area("Ask a custom question:")

            # Display the options for walkthrough
            toggle_walkthrough_ie = toggle_walkthrough_not = False
            if selected_option == "Code Walkthrough":
                st.sidebar.write("Walkthrough details to include:")
                toggle_walkthrough_not = st.sidebar.checkbox("Regular", help=f"Hello Watson! Please use your knowledge as a{code_type} developer and the {code_type} code I send you, to help me with the following tasks or question: I have a very large program but I am going to just show you a part of it. You are going to help me walk through this code and teach me what it is doing and if or how it is interacting with other parts of the larger program.")
                toggle_walkthrough_ie = st.sidebar.checkbox("Nested if/else", help="Do not quote the code in your explanation.")

            # Display checkboxes for selecting the options when "3 Tasks" is selected
            toggle_summarize = toggle_explanation = toggle_interactions = False
            if selected_option == "Code Summarization":
                st.sidebar.write("Summarization details to include:")
                toggle_summarize = st.sidebar.checkbox("Code Overview", help=f"Hello Watson! Please use your knowledge as a {code_type} developer and the {code_type} code I send you, to help me with the following tasks or question: Please summarize what the code is doing.")
                toggle_explanation = st.sidebar.checkbox("Detailed Explanation", help= f"Hello Watson! Please use your knowledge as a {code_type} developer and the {code_type} code I send you, to help me with the following tasks or question: Please explain to me what the code is doing.")
                toggle_interactions = st.sidebar.checkbox("Code Interactions", help= f"Hello Watson! Please use your knowledge as a {code_type} developer and the {code_type} code I send you, to help me with the following tasks or question: If the code is calling or interacting with different parts of the code, tell me how and why.")

            # Update toggle attributes
            self.toggle_walkthrough_ie = toggle_walkthrough_ie
            self.toggle_walkthrough_not = toggle_walkthrough_not
            self.toggle_summarize = toggle_summarize
            self.toggle_explanation = toggle_explanation
            self.toggle_interactions = toggle_interactions
            self.custom_question = custom_question

            # Create a condition to check if at least one checkbox is selected
            at_least_one_checkbox_selected = (toggle_summarize or toggle_explanation or toggle_interactions or toggle_walkthrough_ie or toggle_walkthrough_not or custom_question)

            # SUBMIT (in sidebar)----->    
            # Display the submit button in the sidebar
            submit_button = st.sidebar.button("Submit")
            if submit_button:
                # Check if the "Custom Question" option is selected and if a question is provided
                if selected_option == "Custom Question" and not self.custom_question:
                    st.sidebar.warning("Please type a question before submitting.")
                else:
                    # Check if at least one checkbox is selected for other options
                    at_least_one_checkbox_selected = (
                        self.toggle_summarize or self.toggle_explanation or self.toggle_interactions or
                        self.toggle_walkthrough_ie or self.toggle_walkthrough_not
                    )
                    if selected_option != "Custom Question" and not at_least_one_checkbox_selected:
                        # Display a warning message if none of the checkboxes are checked
                        st.sidebar.warning("Please check at least one checkbox before hitting Submit.")
                    else:
                        self.handle_submit(col2, code_type, context) 
            
        #SAVE (in sidebar)---->
        #Display the save button in the sidebar
        save_button = st.sidebar.button("I like it-Save it!", disabled=not submit_button)
        if save_button:
            self.handle_save()

   #at the end of submit after the api is called and response is displayed...save values in case the user wants to click save
    def store_values_for_save(self, toggle_choice, watson_response):
        # Store values in session state
        st.session_state.toggle_choice = toggle_choice
        st.session_state.watson_response = watson_response

    #when the user clicks save it grabs the values that had been saved by submit
    def retrieve_values_for_save(self):
        # Retrieve stored values
        toggle_choice =st.session_state.toggle_choice
        watson_response = st.session_state.watson_response
        return toggle_choice, watson_response     

    #user clicks save
    def handle_save(self):
        try:
            # Retrieve values for saving
            toggle_choice, watson_response = self.retrieve_values_for_save()         
            self.save_to_word_doc(self.text_area_code, watson_response, toggle_choice)
            st.sidebar.success("Your work has been saved in src/CodeExplain_Output.docx.")
        except Exception as e:
            st.error(f"Error during saving: {str(e)}")

    #user clicks submit
    def handle_submit(self, col2, code_type, context):
        # SUBMIT (in sidebar)---->
        # Check if 'Add Context' is checked and context length exceeds 700 characters
        if not st.session_state.get('context_valid', True):
            st.sidebar.error("Error: Context is too long (max 850 characters).")
            return
        
        # Check the character count of the pasted code
        if self.selected_code == "Paste in my own code":
            code_length = len(self.text_area_code)
            if code_length > 6000:
                st.sidebar.error("Error: Code is too long (max 6000 characters).")
                return

        # Create a placeholder in the sidebar
        status_placeholder = st.sidebar.empty()
        # Initialize the progress bar with "Thinking..."
        progress_text = "Thinking..."
        progress_bar = st.sidebar.progress(0, text=progress_text)

        toggle_choice, watson_response = self.get_toggle_choice(code_type, context)
        # If the user selected "Paste in my own code", get the code from the text area
        if self.selected_code == "Paste in my own code":
            self.text_area_code = st.session_state.pasted_code

        # Update the progress bar to "In Progress"
        progress_bar.progress(50, text="Responding...")

        # Simulate a long-running process (replace this with your actual processing logic)
        for percent_complete in range(51, 101):
            # Update the progress bar
            progress_bar.progress(percent_complete, text="Responding...")

            # Simulate some processing time
            time.sleep(0.1)

        # Update status after return of watson response
        progress_bar.progress(percent_complete, text= "Done.")
        # Display watson response
        with col2:
            st.text_area("Response from Watsonx.ai:", value=watson_response, height=1020)

        # Store values for saving
        self.store_values_for_save(toggle_choice, watson_response)

        # Reset session state variables
        st.session_state.context_valid = True

        # Return toggle_choice and watson_response
        return toggle_choice, watson_response

    # Extract the logic for generating toggle_choice to a separate method
    def get_toggle_choice(self, code_type, context):
        # Check if watson_response has already been fetched
        if hasattr(self, 'watson_response') and self.watson_response is not None:
            return self.toggle_choice, self.watson_response

        # If not, fetch them and store them
        toggle_choice = ", ".join(filter(None, [
            "Please summarize what the code is doing." if self.toggle_summarize else "",
            "Please explain to me what the code is doing." if self.toggle_explanation else "",
            "If the code is calling or interacting with different parts of the code, tell me how and why." if self.toggle_interactions else "",
            "I have a very large program but I am going to just show you a part of it. You are going to help me walk through this code and teach me what it is doing and if or how it is interacting with other parts of the larger program." if self.toggle_walkthrough_not else "",
            "Do not quote the code in your explanation." if self.toggle_walkthrough_ie else "",
            self.custom_question
        ]))
        watson_response = call_watsonx(toggle_choice, self.text_area_code, code_type, context)

        return toggle_choice, watson_response
    
    def run(self):
        #  Inject CSS for styling
        self.inject_custom_css()
        # Display the logo in the logo_column
        logo_image = "src/ibmLogo.jpg"
        # For windows use this
        #logo_image = "src\ibmLogo.jpg"
        st.image(logo_image, width = 350)

        # Display the title in the title_column
        st.subheader("CodeExplain with IBM watsonx")

        # File upload in the sidebar can switch cbl to cob depending on file type
        uploaded_file = st.sidebar.file_uploader("Upload Code File", type=["cbl", "asm", "py", "java","bas"])
        
        if uploaded_file:
            
            try:
                stringio = StringIO(uploaded_file.getvalue().decode('utf-8'))
                content = stringio.read()
                chunks, code_type = self.process_file_by_type(content, uploaded_file.name)
                self.update_ui_after_file_upload(chunks, code_type)
            except FileNotFoundError:
                st.error("File not found.")
            except Exception as e:
                st.error(f"Error reading the file: {str(e)}")
    
    @staticmethod
    def inject_custom_css():
        """Injects custom CSS for styling specific Streamlit elements."""
        custom_css = """
        <style>
            .stButton > button {
                background-color: #fd7e14;
                color: white;
            }
            .stFileUploader > div > div > button {
                background-color: #fd7e14;
                color: white;
            }
        </style>
        """
        st.markdown(custom_css, unsafe_allow_html=True)

    def save_to_word_doc(self, code, watson_response, prompt):
        doc_filename = "CodeExplain_Output.docx"

        # Check if the document file already exists
        if os.path.isfile(doc_filename):
            # If it exists, open the document
            doc = Document(doc_filename)
        else:
            # If it doesn't exist, create a new document
            doc = Document()

        # Add a new title with the current date and time
        doc.add_heading(f"CodeExplain Output - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=0)

        # Add input code
        doc.add_heading("Prompt", level=1)
        doc.add_paragraph(prompt)

        # Add Watson's response
        doc.add_heading("Watson's Response", level=1)
        doc.add_paragraph(watson_response)

        # Add input code
        doc.add_heading("Input Code", level=1)
        doc.add_paragraph(code)

        # Save the document
        doc.save(doc_filename)

if __name__ == "__main__":
    app = StreamlitApp()
    app.run()