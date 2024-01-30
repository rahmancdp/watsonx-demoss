import streamlit as st

from ibm_watsonx_ai.foundation_models.utils.enums import ModelTypes
from ibm_watsonx_ai.foundation_models import Model
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from ibm_watsonx_ai.foundation_models.extensions.langchain import WatsonxLLM

import tempfile
import pathlib

from assemblerCode import split_assem_code, clean_assem_code
from javaCode import split_java_code
from cobolCode import split_smaller
from docx import Document

import os
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("API_KEY", None)
project_id = os.getenv("PROJECT_ID", None)

creds = {
    "url"    : "https://us-south.ml.cloud.ibm.com",
    "apikey" : api_key
}

params = {
    GenParams.DECODING_METHOD:"greedy",
    GenParams.MAX_NEW_TOKENS:4000,
    GenParams.MIN_NEW_TOKENS:1,
    GenParams.TOP_K:50,
    GenParams.TOP_P:1,
    GenParams.STOP_SEQUENCES:['```'],
}

import ssl

try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context


# llmllama = Model("bigcode/starcoder",creds, params, project_id)
llmllama = Model("meta-llama/llama-2-70b-chat",creds, params, project_id)

def buildpromptforconvertpseudocode(pseudocode):
    return f"""Assume you are an expert both in java. 
    The code below is in Pseudo Code. 
    Please re-write in the java language, keeping the same logic.  

### Pseudo Code:
```{pseudocode}```

Please re-write the code above in java language. 

Here are the requirements:

1. Indicate the parent class when you define a public class.
2. The original code can be a snippet from complete code base. 
3. If some classes are not provided in the given code snippet, import them from external packages. 
4. Mark with comments if things may not be correct through direct conversion or things can be missing.
5. Make sure the code concise, brief and elegant.
6. Import java classes as few as possible. If necessary use ".*" to represent multiple classes.

### java code:
```java"""

def buildpromptforconvert(code):
    return f"""Assume you are an expert both in COBOL and java. The code below is in COBOL. Please re-write in the java language, keeping the same logic. Replace functions from COBOL with the equivalent in java that makes the most sense. 

### COBOL Code:
```{code}```

Please re-write the code above in java language. 

Here are the requirements:

1. Indicate the parent class when you define a public class.
2. The original code can be a snippet from complete code base. 
3. If some classes are not provided in the given code snippet, import them from external packages. 
4. Mark with comments if things may not be correct through direct conversion or things can be missing.
5. Make sure the code concise, brief and elegant.
6. Import java classes as few as possible. If necessary use ".*" to represent multiple classes.

### java code:
```java"""

def buildpromptforDDD(code):
    return f"""Assume you are an expert on java. The code below is in java. 
    Please re-write in the java language with DDD style, keeping the same logic. 

### java Code:
```{code}```

Please re-write the code above in java language with DDD style. 
- add function description with input/output
- add comment on complex logic block
- function name in pattern 'func_HSBC'+business+function name
- variable name in pattern 'var_'+business+variable name

### java code with DDD style:
```java"""

def buildpromptfordoc(code):
    return f"""[INST]
please generate documentation for the java source code in backquoted
<<SYS>>
`{code}`
<<SYS>>
[/INST]
documentation:"""

def buildpromptfortestcase(code):
    return f"""[INST]
please generate testcase in junit for the java source code in backquoted
- test for all conditional case
- test for all possible output
- test for all possible input
- test for exception catching
<<SYS>>
`{code}`
<<SYS>>
[/INST]
testcase:"""

def buildpromptforexplaination(code):
    return f"""[INST]you are expert on COBOL programming,
please explain the programming logic for the COBOL source code in backquoted.
- the explaination should show the logic of the code.
- the explaination should be well structure, easy to understand.
- the explaination should be ready for modernize to java language.
<<SYS>>
`{code}`
<<SYS>>
[/INST]
explaination:"""

def genexplaination(incode):
    #chunking
    chunks = split_smaller(incode)

    prompts = []
    for chunk in chunks:
        print(len(chunk))
        prompts += [buildpromptforexplaination(chunk)]
    code = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            code += response
    print(code)
    return code

def convertcoboltojava(incode):
    #chunking
    chunks = split_smaller(incode)

    prompts = []
    for chunk in chunks:
        print(len(chunk))
        prompts += [buildpromptforconvert(chunk)]
    code = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            code += response
    return code

def convertpseudocodetojava(incode):
    #chunking
    chunks = split_smaller(incode)

    prompts = []
    for chunk in chunks:
        print(len(chunk))
        prompts += [buildpromptforconvertpseudocode(chunk)]
    code = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            code += response
    return code

def genDDD(incode):
    #chunking
    chunk_size = len(incode)//4
    chunks = []
    for i in range(0, len(incode), chunk_size):
        chunks += [incode[i:i+chunk_size]]

    prompts = []
    for chunk in chunks:
        prompts += [buildpromptforDDD(chunk)]
    code = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            code += response
    return code.replace("```","")

def gendoc(incode):
    #chunking
    chunk_size = len(incode)//4
    chunks = []
    for i in range(0, len(incode), chunk_size):
        chunks += [incode[i:i+chunk_size]]

    prompts = []
    for chunk in chunks:
        prompts += [buildpromptfordoc(chunk)]
    doc = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            doc += response
    return doc

def gentestcase(incode):
    #chunking
    chunk_size = len(incode)//4
    chunks = []
    for i in range(0, len(incode), chunk_size):
        chunks += [incode[i:i+chunk_size]]

    prompts = []
    for chunk in chunks:
        prompts += [buildpromptfortestcase(chunk)]
    testcase = ""
    for response in llmllama.generate_text(prompts):
        if response is not None:
            testcase += response
    return testcase

def querycodeanddoc(informations, question):
    prompt = f"""
    answer the question in 5 sentences base on the informations:
    informations:
    {informations}
    question:
    {question}
    answer in point form:"""

    prompts = [prompt]
    answer = ""
    for response in llmllama.generate_text(prompts):
        answer += response
    return answer

def process_file_by_type(file, file_path):
        # Determine file type based on extension
        file_extension = file_path.split('.')[-1].lower()

        if file_extension == 'cbl':
            code_type = "cobol"
            #remove the prior to procedure division section
            file_content=split_off_procedure(file)
            #clean up the cobol code
            pretty_cobol =process_cobol_file(file_content)
            #split cobol into watson sized chunks
            cobol_chunks = split_cbl_code(pretty_cobol)
            return cobol_chunks, code_type
        elif file_extension == 'asm':
            code_type = "assembler"
            pretty_assem = clean_assem_code(file)
            assembler_chunks =split_assem_code(pretty_assem)
            return assembler_chunks, code_type
        elif file_extension == 'java':
            code_type = "java"
            java_chunks = split_java_code(file)
            return java_chunks, code_type  
        else:
            print(f"Unsupported file type: {file_extension}")
            return []

def save_to_word_doc(code, watson_response, prompt):
    doc_filename = "CodeExplain_Output.docx"

    # Check if the document file already exists
    if os.path.isfile(doc_filename):
        # If it exists, open the document
        doc = Document(doc_filename)
    else:
        # If it doesn't exist, create a new document
        doc = Document()

    # Add a new title with the current date and time
    doc.add_heading(f"CodeExplain Output - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=0)

    # Add input code
    doc.add_heading("Prompt", level=1)
    doc.add_paragraph(prompt)

    # Add Watson's response
    doc.add_heading("Watson's Response", level=1)
    doc.add_paragraph(watson_response)

    # Add input code
    doc.add_heading("Input Code", level=1)
    doc.add_paragraph(code)

    # Save the document
    doc.save(doc_filename)

temp_dir = tempfile.TemporaryDirectory()

st.set_page_config(layout="wide")
st.subheader("modernization powered by watsonx")

st.markdown("""
    <style>
        .reportview-container {
            margin-top: -2em;
        }
        #MainMenu {visibility: hidden;}
        .stDeployButton {display:none;}
        footer {visibility: hidden;}
        #stDecoration {display:none;}
    </style>
""", unsafe_allow_html=True)

hide_streamlit_style = """
                <style>
                div[data-testid="stToolbar"] {
                visibility: hidden;
                height: 0%;
                position: fixed;
                }
                div[data-testid="stDecoration"] {
                visibility: hidden;
                height: 0%;
                position: fixed;
                }
                div[data-testid="stStatusWidget"] {
                visibility: hidden;
                height: 0%;
                position: fixed;
                }
                #MainMenu {
                visibility: hidden;
                height: 0%;
                }
                header {
                visibility: hidden;
                height: 0%;
                }
                footer {
                visibility: hidden;
                height: 0%;
                }
                </style>
                """
st.markdown(hide_streamlit_style, unsafe_allow_html=True) 

if "cobolin" not in st.session_state:
    st.session_state.cobolin = ""
if "javaout" not in st.session_state:
    st.session_state.javaout = ""
if "javawithddd" not in st.session_state:
    st.session_state.javawithddd = ""
if "outdoc" not in st.session_state:
    st.session_state.outdoc = ""
if "explaination" not in st.session_state:
    st.session_state.explaination = ""
if "testcase" not in st.session_state:
    st.session_state.testcase = ""

uploaded_file = st.file_uploader(label="upload a source code file",type=['cbl','asm'])
if uploaded_file is not None:
    st.write("filename:", uploaded_file.name)
    with open(os.path.join(pathlib.Path(temp_dir.name),uploaded_file.name),"wb") as f:
        f.write(uploaded_file.getbuffer())
        st.session_state.cobolin = uploaded_file.read().decode()

tabexplain, tabconvert, tabDDD, tabdoc, tabtestcase = st.tabs(['explain','java','DDD','Documentation','Test Case'])

with tabexplain:
    st.subheader("explain the COBOL")
    colsource, coloutcome = st.columns(2)

    with colsource:
        st.code(st.session_state.cobolin,language="cobol")

    with coloutcome:
        btnExplain = st.button("explain")
        if btnExplain:
            # st.write(sourcecode['text'])
            with st.spinner(text="In progress...", cache=False):
                st.session_state.explaination = genexplaination(st.session_state.cobolin)
        st.markdown(st.session_state.explaination)

with tabconvert:
    st.subheader("convert Pseudo Code to Java")
    colsource, coloutcome = st.columns(2)
    with colsource:
        st.markdown(st.session_state.explaination)
    with coloutcome:
        btnJava = st.button("convert to java")
        if btnJava:
            # st.write(sourcecode['text'])
            with st.spinner(text="In progress...", cache=False):
                st.session_state.javaout = convertpseudocodetojava(st.session_state.explaination)
        st.code(st.session_state.javaout,language="java")

with tabDDD:
    st.subheader("convert Java to DDD style")
    colsource, coloutcome = st.columns(2)
    with colsource:
        st.code(st.session_state.javaout,language="java")
    with coloutcome:
        btnDDD = st.button("convert to DDD")
        if btnDDD:
            # st.write(sourcecode['text'])
            with st.spinner(text="In progress...", cache=False):
                st.session_state.javawithddd = genDDD(st.session_state.javaout)
        st.code(st.session_state.javawithddd,language="java")

with tabdoc:
    st.subheader("Generate Documentation")
    colsource, coloutcome = st.columns(2)
    with colsource:
        st.code(st.session_state.javawithddd,language="java")
    with coloutcome:
        btnDoc = st.button("generate documentation")
        if btnDoc:
            with st.spinner(text="In progress...", cache=False):
                st.session_state.outdoc = gendoc(st.session_state.javawithddd)
        st.markdown(st.session_state.outdoc)

with tabtestcase:
    st.subheader("Generate Test Case")
    colsource, coloutcome = st.columns(2)
    with colsource:
        st.code(st.session_state.javawithddd,language="java")
    with coloutcome:
        btnTestCase = st.button("generate testcase")
        if btnTestCase:
            with st.spinner(text="In progress...", cache=False):
                st.session_state.testcase = gentestcase(st.session_state.javawithddd)
        st.markdown(st.session_state.testcase)